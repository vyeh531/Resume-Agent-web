from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.comments import Comment
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation


OUT_DIR = Path("outputs/advice_data_brief")
DOCX_PATH = OUT_DIR / "Advice_Data_Handoff_Guide.docx"
XLSX_PATH = OUT_DIR / "Advice_Data_Collection_Template.xlsx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"
WHITE = "FFFFFF"


ROLE_FAMILIES = [
    "it_support",
    "data_engineer",
    "communications_pr",
    "legal_compliance",
    "life_sciences",
    "industrial_quality",
    "civil_construction",
    "actuarial",
    "procurement",
    "hospitality_events",
    "journalism_media",
    "policy_public_sector",
    "research_academic",
    "social_services",
]

PROBLEM_TAGS = [
    "missing_exact_job_title",
    "weak_summary_role_alignment",
    "weak_target_role_alignment",
    "generic_resume_positioning",
    "resume_not_tailored_to_jd",
    "low_jd_keyword_match",
    "missing_priority_keywords",
    "low_hard_skill_match",
    "weak_experience_keyword_evidence",
    "keywords_only_in_skills",
    "low_measurable_results",
    "missing_github_link",
    "missing_portfolio",
]

TASK_BUCKETS = ["positioning", "keyword_evidence", "artifact_links"]
SENIORITY = ["student", "new_grad", "entry_level", "early_career", "career_switcher", "experienced"]
STATUS = ["draft", "needs_review", "approved", "needs_rewrite"]


FIELD_DEFS = [
    ("role_family", "Required", "Use one controlled role family, e.g. data_engineer or legal_compliance."),
    ("target_role_examples", "Required", "Concrete job titles this advice applies to."),
    ("task_bucket", "Required", "One of positioning, keyword_evidence, artifact_links."),
    ("problem_tags", "Required", "1-3 tags from the reference list, comma-separated."),
    ("advice_card_title", "Required", "Short card title shown to the user."),
    ("user_problem_summary", "Required", "Plain-language diagnosis of the student's resume issue."),
    ("mentor_insight", "Required", "Why this matters from a mentor/industry perspective."),
    ("action_summary", "Required", "Specific resume edit action the student can take."),
    ("example_before", "Recommended", "Weak or generic bullet/summary example."),
    ("example_after", "Recommended", "Improved example with role language and evidence."),
    ("keywords", "Required", "Role-specific keywords that make retrieval work."),
    ("seniority", "Required", "Use student/new_grad/entry_level unless advice is for experienced candidates."),
    ("source_or_mentor", "Recommended", "Who provided it, or 'internal template' if no specific mentor."),
    ("quality_check_notes", "Optional", "Reviewer comments, caveats, or missing context."),
]


EXAMPLE_ROWS = [
    {
        "role_family": "it_support",
        "target_role_examples": "Help Desk Technician, IT Support Specialist, Systems Administrator",
        "task_bucket": "keyword_evidence",
        "problem_tags": "low_hard_skill_match,weak_experience_keyword_evidence,keywords_only_in_skills",
        "advice_card_title": "把 Help Desk 經驗寫成 troubleshooting 證據",
        "user_problem_summary": "履歷只寫協助處理 IT 問題，但沒有說明處理的是什麼系統、什麼 ticket、什麼工具。",
        "mentor_insight": "IT Support 類崗位很看重候選人是否能穩定處理真實使用者問題；只寫協助會讓招聘方看不出技術範圍和獨立處理能力。",
        "action_summary": "把經歷改成 ticket 類型 + 使用工具 + 解決結果的格式，補上 ServiceNow、Windows/Mac、Active Directory、VPN、printer/network troubleshooting 等真實做過的內容。",
        "example_before": "Assisted with IT support requests.",
        "example_after": "Resolved 20+ weekly support tickets for Windows login, VPN access, printer setup, and software installation issues using ServiceNow, reducing repeated escalations to senior IT staff.",
        "keywords": "ServiceNow, ticketing, troubleshooting, Windows, Active Directory, VPN, help desk",
        "seniority": "student,new_grad,entry_level",
        "source_or_mentor": "internal template",
    },
    {
        "role_family": "data_engineer",
        "target_role_examples": "Data Engineer, Analytics Engineer, ETL Developer",
        "task_bucket": "keyword_evidence",
        "problem_tags": "low_jd_keyword_match,missing_priority_keywords,weak_experience_keyword_evidence",
        "advice_card_title": "不要只寫 SQL/Python，要寫 pipeline 證據",
        "user_problem_summary": "履歷有 SQL 和 Python，但沒有把資料來源、ETL 流程、排程工具或 data warehouse 產出寫清楚。",
        "mentor_insight": "Data Engineer 篩選時會看你是否理解資料流和 production pipeline，不只是會查資料或做分析。",
        "action_summary": "在 project bullet 裡補資料來源、處理步驟、工具和產出，例如 Airflow/Spark/dbt/Snowflake/BigQuery，以及 pipeline 的更新頻率或資料量。",
        "example_before": "Built data project using Python and SQL.",
        "example_after": "Built an Airflow ETL pipeline that ingested daily sales CSVs into Snowflake, transformed raw tables with dbt, and produced dashboard-ready marts for weekly KPI reporting.",
        "keywords": "ETL, data pipeline, Airflow, Spark, dbt, Snowflake, data warehouse, orchestration",
        "seniority": "student,new_grad,entry_level",
        "source_or_mentor": "internal template",
    },
    {
        "role_family": "communications_pr",
        "target_role_examples": "PR Specialist, Communications Associate, Content Strategist",
        "task_bucket": "artifact_links",
        "problem_tags": "missing_portfolio,weak_experience_keyword_evidence",
        "advice_card_title": "補上可點擊的 writing sample 或 campaign portfolio",
        "user_problem_summary": "履歷提到寫作、社群或 campaign，但沒有提供作品連結，招聘方難以判斷實際輸出品質。",
        "mentor_insight": "PR/Comms 類崗位很看重可展示的文字、媒體稿、社群內容或 campaign 產出；作品比泛泛描述更能建立信任。",
        "action_summary": "在姓名下方或 Projects/Portfolio section 加入 2-4 個精選 writing samples/campaign links，並在 bullet 裡說明你的角色、受眾、渠道和結果。",
        "example_before": "Created social media content for school organization.",
        "example_after": "Wrote and scheduled 25+ Instagram/LinkedIn posts for a campus campaign, increasing event sign-ups by 32%; portfolio: [link].",
        "keywords": "writing sample, press release, media relations, campaign, content calendar, social media, portfolio",
        "seniority": "student,new_grad,entry_level",
        "source_or_mentor": "internal template",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Advice Data 補充任務交接指南")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("給不熟悉 project 的同事：如何補可進資料庫的導師履歷建議素材")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_border(cell, "CBD5E1")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = hdr[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(value))
    doc.add_paragraph()
    return table


def build_docx():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    add_callout(
        doc,
        "一句話任務",
        "我們要補的不是一般履歷建議，而是「某個崗位 + 某個履歷問題 + 具體怎麼改」的導師 advice data。每條都要能直接放進系統做檢索和展示。",
    )

    doc.add_heading("1. Project 背景：你需要知道的就這些", level=1)
    add_bullets(doc, [
        "系統會分析學生的履歷和 JD，判斷他投什麼崗位、履歷缺什麼。",
        "系統再從導師建議庫找最匹配的 advice，展示給學生。",
        "目前一些崗位方向缺少 role-specific advice，所以會抓到太泛或不完全適合的建議。",
        "你的任務是補一批標準化 advice，讓系統未來可以更準確匹配。",
    ])

    doc.add_heading("2. 優先補哪些崗位", level=1)
    rows = [
        ("P0", "it_support", "Help Desk, IT Support, System Administrator"),
        ("P0", "data_engineer", "Data Engineer, Analytics Engineer, ETL Developer"),
        ("P0", "communications_pr", "PR, Communications, Content, Journalism"),
        ("P0", "legal_compliance", "Compliance, Legal Assistant, Policy Analyst"),
        ("P0", "life_sciences", "Biotech, Research Assistant, Clinical Research"),
        ("P1", "industrial_quality / civil_construction", "Quality, Manufacturing, Civil, Construction"),
        ("P1", "actuarial / procurement / hospitality / social_services", "Actuarial, Procurement, Hospitality, Social Services"),
    ]
    add_table(doc, ["Priority", "Role family", "Examples"], rows, [0.8, 2.1, 3.6])

    doc.add_heading("3. 每個崗位先補兩類 advice", level=1)
    add_table(
        doc,
        ["Task bucket", "要解決的問題", "典型 tags"],
        [
            ("positioning", "履歷定位不清、summary 沒有目標職位、經歷沒有翻成該崗位語言。", "missing_exact_job_title, weak_summary_role_alignment, weak_target_role_alignment, resume_not_tailored_to_jd"),
            ("keyword_evidence", "缺 JD 技能詞，或技能只放 Skills，沒有放進經歷 bullet。", "low_jd_keyword_match, missing_priority_keywords, low_hard_skill_match, weak_experience_keyword_evidence, keywords_only_in_skills"),
            ("artifact_links", "需要 GitHub、portfolio、writing sample、case study，但履歷沒有放或放得不清楚。", "missing_github_link, missing_portfolio"),
        ],
        [1.35, 3.0, 2.15],
    )

    doc.add_heading("4. 每條 advice 必須包含什麼", level=1)
    add_table(
        doc,
        ["欄位", "是否必填", "怎麼填"],
        FIELD_DEFS,
        [1.65, 0.85, 4.0],
    )

    doc.add_heading("5. 好 advice 和不合格 advice 的差別", level=1)
    add_table(
        doc,
        ["不合格：太泛", "合格：可進資料庫"],
        [
            (
                "根據 JD 修改履歷，突出相關技能。",
                "針對 Data Engineer，不能只寫 Python/SQL。要把 pipeline、ETL、Airflow、Spark、Snowflake、data warehouse 放進實際 project bullet，並說明資料來源、處理流程和產出。",
            ),
            (
                "多寫一些自己的工作成果。",
                "針對 IT Support，把「協助 IT 問題」改成 ticket 類型、工具和解決結果，例如 ServiceNow、VPN、Active Directory、Windows login、printer setup。",
            ),
            (
                "作品集很重要，建議補上。",
                "針對 PR/Communications，補 2-4 個 writing sample 或 campaign links，並在 bullet 說明受眾、渠道、你的角色和結果。",
            ),
        ],
        [3.0, 3.5],
    )

    doc.add_heading("6. 建議工作量", level=1)
    add_bullets(doc, [
        "每個 role family 先補 8-10 條。",
        "其中 3 條 positioning advice，5 條 keyword/evidence advice。",
        "如果是 tech/data/design/PR/journalism，再補 1-2 條 GitHub/portfolio/writing sample 類 advice。",
        "第一輪先補 5 個 P0 role family，大約 40-50 條，就能明顯改善覆蓋。",
    ])

    doc.add_heading("7. 填寫流程", level=1)
    for idx, step in enumerate([
        "先在 Excel 的 Role Priorities sheet 選一個 role_family。",
        "到 Data Entry sheet 新增資料列。",
        "用 dropdown 選 task_bucket、seniority、status。",
        "problem_tags 填 1-3 個，用逗號隔開。",
        "寫 user_problem_summary、mentor_insight、action_summary。",
        "盡量補 example_before / example_after，讓內容更可用。",
        "最後自查：是否足夠具體？是否真的適用這個崗位？是否有可檢索 keywords？",
    ], start=1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.add_run(step)

    doc.add_heading("8. 快速品質檢查", level=1)
    add_bullets(doc, [
        "不要只寫泛泛的「根據 JD 修改履歷」。要寫出該崗位常見技能、工具、交付物。",
        "不要讓不同崗位混線：Legal advice 不要寫成 Data advice；IT Support 不要寫成 SWE advice。",
        "每條 action_summary 要能讓學生照著改一段 summary 或 bullet。",
        "如果沒有確定的導師來源，source_or_mentor 先填 internal template，不要假裝是某位導師原話。",
    ])

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Advice data handoff guide")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("777777")

    doc.save(DOCX_PATH)


def style_header(row):
    for cell in row:
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.font = Font(color=WHITE, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=Side(style="thin", color=BORDER))


def apply_table_style(ws, max_row, max_col):
    thin = Side(style="thin", color=BORDER)
    for row in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    style_header(ws[1])
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(max_col)}{max_row}"


def build_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Data Entry"

    headers = [
        "status",
        "role_family",
        "target_role_examples",
        "task_bucket",
        "problem_tags",
        "advice_card_title",
        "user_problem_summary",
        "mentor_insight",
        "action_summary",
        "example_before",
        "example_after",
        "keywords",
        "seniority",
        "source_or_mentor",
        "quality_check_notes",
    ]
    ws.append(headers)
    for row in EXAMPLE_ROWS:
        ws.append([
            "draft",
            row["role_family"],
            row["target_role_examples"],
            row["task_bucket"],
            row["problem_tags"],
            row["advice_card_title"],
            row["user_problem_summary"],
            row["mentor_insight"],
            row["action_summary"],
            row["example_before"],
            row["example_after"],
            row["keywords"],
            row["seniority"],
            row["source_or_mentor"],
            "",
        ])
    for _ in range(80):
        ws.append(["draft"] + [""] * (len(headers) - 1))

    widths = [14, 22, 34, 18, 38, 30, 48, 54, 58, 42, 58, 42, 22, 24, 34]
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    for row in range(1, ws.max_row + 1):
        ws.row_dimensions[row].height = 60 if row > 1 else 34
    apply_table_style(ws, ws.max_row, len(headers))

    comments = {
        "A1": "Use draft / needs_review / approved / needs_rewrite.",
        "B1": "Choose one role_family from Reference Lists.",
        "D1": "positioning, keyword_evidence, or artifact_links.",
        "E1": "Use 1-3 problem tags, comma-separated.",
        "F1": "Short user-facing title.",
        "G1": "Plain-language diagnosis.",
        "H1": "Why this matters from mentor/industry perspective.",
        "I1": "Specific action the student can take on the resume.",
        "L1": "Role-specific keywords for retrieval.",
        "M1": "student,new_grad,entry_level etc.",
    }
    for cell, text in comments.items():
        ws[cell].comment = Comment(text, "Codex")

    role_sheet = wb.create_sheet("Role Priorities")
    role_sheet.append(["priority", "role_family", "target_role_examples", "task_buckets_to_fill", "suggested_count", "notes"])
    priority_rows = [
        ("P0", "it_support", "Help Desk Technician, IT Support Specialist, Systems Administrator", "positioning, keyword_evidence, artifact_links", 8, "Do not turn into SWE advice; focus on tickets, troubleshooting, systems, user support."),
        ("P0", "data_engineer", "Data Engineer, Analytics Engineer, ETL Developer", "positioning, keyword_evidence, artifact_links", 8, "Focus on pipeline, ETL, warehouse, orchestration, data quality."),
        ("P0", "communications_pr", "PR Specialist, Communications Associate, Content Strategist, Journalist", "positioning, keyword_evidence, artifact_links", 8, "Focus on writing samples, campaigns, audience, channels, measurable reach."),
        ("P0", "legal_compliance", "Compliance Analyst, Legal Assistant, Policy Analyst", "positioning, keyword_evidence", 8, "Needs domain-specific language; do not reuse finance/data advice blindly."),
        ("P0", "life_sciences", "Research Assistant, Clinical Research Associate, Biotech Associate", "positioning, keyword_evidence", 8, "Translate lab/research/clinical details into industry resume evidence."),
        ("P1", "industrial_quality", "Quality Engineer, Manufacturing Quality Analyst", "positioning, keyword_evidence", 6, "Focus on QA, root cause, process improvement, CAPA, inspection."),
        ("P1", "civil_construction", "Civil Engineer, Construction Coordinator, Construction Manager", "positioning, keyword_evidence", 6, "Focus on site, schedule, safety, AutoCAD, contractor coordination."),
        ("P1", "actuarial", "Actuarial Analyst, Actuarial Intern", "positioning, keyword_evidence", 6, "Focus on exams, modeling, insurance, risk, Excel/R/Python."),
        ("P1", "procurement", "Procurement Specialist, Buyer, Purchasing Agent", "positioning, keyword_evidence", 6, "Focus on vendors, sourcing, PO, negotiation, cost savings."),
        ("P1", "hospitality_events", "Event Planner, Hospitality Coordinator", "positioning, keyword_evidence", 6, "Focus on vendors, guest experience, logistics, budget, coordination."),
        ("P1", "social_services", "Social Worker, Counselor, Case Manager", "positioning, keyword_evidence", 6, "Focus on case management, assessment, client communication, care plans."),
    ]
    for row in priority_rows:
        role_sheet.append(row)
    for idx, width in enumerate([12, 24, 48, 36, 16, 72], start=1):
        role_sheet.column_dimensions[get_column_letter(idx)].width = width
    for row in range(1, role_sheet.max_row + 1):
        role_sheet.row_dimensions[row].height = 44 if row > 1 else 30
    apply_table_style(role_sheet, role_sheet.max_row, 6)

    ref = wb.create_sheet("Reference Lists")
    ref.append(["role_family", "problem_tags", "task_bucket", "seniority", "status"])
    max_len = max(len(ROLE_FAMILIES), len(PROBLEM_TAGS), len(TASK_BUCKETS), len(SENIORITY), len(STATUS))
    for i in range(max_len):
        ref.append([
            ROLE_FAMILIES[i] if i < len(ROLE_FAMILIES) else "",
            PROBLEM_TAGS[i] if i < len(PROBLEM_TAGS) else "",
            TASK_BUCKETS[i] if i < len(TASK_BUCKETS) else "",
            SENIORITY[i] if i < len(SENIORITY) else "",
            STATUS[i] if i < len(STATUS) else "",
        ])
    for idx, width in enumerate([28, 38, 22, 22, 18], start=1):
        ref.column_dimensions[get_column_letter(idx)].width = width
    apply_table_style(ref, ref.max_row, 5)

    guide = wb.create_sheet("Field Guide")
    guide.append(["field", "required", "instruction"])
    for row in FIELD_DEFS:
        guide.append(row)
    for idx, width in enumerate([24, 16, 86], start=1):
        guide.column_dimensions[get_column_letter(idx)].width = width
    for row in range(1, guide.max_row + 1):
        guide.row_dimensions[row].height = 38 if row > 1 else 30
    apply_table_style(guide, guide.max_row, 3)

    # Data validation on Data Entry.
    def add_list_validation(col, formula):
        dv = DataValidation(type="list", formula1=formula, allow_blank=True)
        ws.add_data_validation(dv)
        dv.add(f"{col}2:{col}200")

    add_list_validation("A", "'Reference Lists'!$E$2:$E$6")
    add_list_validation("B", "'Reference Lists'!$A$2:$A$30")
    add_list_validation("D", "'Reference Lists'!$C$2:$C$10")
    add_list_validation("M", "'Reference Lists'!$D$2:$D$12")

    required_cols = ["B", "D", "E", "F", "G", "H", "I", "L", "M"]
    for col in required_cols:
        ws.conditional_formatting.add(
            f"{col}2:{col}200",
            FormulaRule(formula=[f'LEN(TRIM({col}2))=0'], fill=PatternFill("solid", fgColor="FFF2CC")),
        )

    # Hide reference sheet gridlines and keep sheet order user-friendly.
    for sheet in wb.worksheets:
        sheet.sheet_view.showGridLines = False
    wb.save(XLSX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_xlsx()
    print(DOCX_PATH.resolve())
    print(XLSX_PATH.resolve())


if __name__ == "__main__":
    main()
